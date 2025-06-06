from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

from src.domain import Table, ColumnFormat


class ApaWordTableFormatter:

    def __init__(self, single_column_width_inches: float, double_column_width_inches: float, font: str, font_size: int):
        self.single_column_width_inches = single_column_width_inches
        self.double_column_width_inches = double_column_width_inches
        self.font = font
        self.font_size = font_size

    def create(
            self,
            table: Table,
            file_name: Path,
            column_format: ColumnFormat = ColumnFormat.DOUBLE
    ):
        # Create a new Word document
        doc = Document()

        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add the table title in italics, same font and style as the table
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(table.title)
        run.italic = True
        run.font.name = self.font
        run.font.size = Pt(self.font_size)
        paragraph.alignment = 0  # Left align
        paragraph.paragraph_format.left_indent = Inches(0)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.right_indent = Inches(0)

        # Create the table in the Word document
        num_rows = len(table.rows) + 1  # +1 for the header
        num_cols = len(table.header)
        doc_table = doc.add_table(rows=num_rows, cols=num_cols)

        # Set table width according to column format
        if column_format == ColumnFormat.SINGLE:
            table_width = Inches(3)
        else:
            table_width = Inches(6.5)

        doc_table.autofit = True
        doc_table.allow_autofit = True
        doc_table.preferred_width = table_width

        # Adjust column widths equally
        for row in doc_table.rows:
            for cell in row.cells:
                cell.width = table_width / num_cols

        # Remove all borders
        tbl = doc_table._tbl
        tblBorders = tbl.xpath('.//w:tblBorders')
        if tblBorders:
            tblBorders[0].getparent().remove(tblBorders[0])

        # Populate the header
        for j, header_text in enumerate(table.header):
            cell = doc_table.cell(0, j)
            cell.text = str(header_text)
            # Set cell formatting
            for paragraph in cell.paragraphs:
                run = paragraph.runs[0]
                font = run.font
                font.name = self.font
                font.size = Pt(self.font_size)
                paragraph.alignment = 1  # Center align
            # Add top and bottom borders to header cells
            self._set_cell_boarder(cell, top={'sz': '12', 'val': 'single'}, bottom={'sz': '6', 'val': 'single'})

        # Populate the rows
        for i, row in enumerate(table.rows, start=1):
            for j, value in enumerate(row):
                cell = doc_table.cell(i, j)
                cell.text = str(value)
                # Set cell formatting
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    font = run.font
                    font.name = self.font
                    font.size = Pt(self.font_size)
                    paragraph.alignment = 1  # Center align
                # Add bottom border to the last row cells
                if i == num_rows - 1:
                    self._set_cell_boarder(cell, bottom={'sz': '12', 'val': 'single'})

        complete_path = file_name.with_suffix('.docx')
        doc.save(str(complete_path))

    @staticmethod
    def _set_cell_boarder(cell, **kwargs):
        """
        Set cell's border
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Check for tag existence, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # List over all available tags
        for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'w:{edge}'

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                for key in ('sz', 'val', 'color', 'space'):
                    if key in edge_data:
                        element.set(qn(f'w:{key}'), str(edge_data[key]))
